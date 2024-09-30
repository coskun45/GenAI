import textwrap
from IPython.display import display, Markdown
def print_formatted_md(md_text, max_line_length=80):
    lines=md_text.splitlines()
    wrapped_lines=[]
    
    for line in lines:
        #handle list items and paragraphs with bold text
        if line.startswith(('* ', '- ', '> ', '**')):
            prefix=""
            if line.startswith('**'):
                prefix = '**'
                line = line[2:]
            elif line.startswith('- '):
                prefix = '-'
                line = line[2:]
            elif line.startswith('> '):
                prefix='> '
                line = line[2:]
            
            wrapped = textwrap.wrap(line, width=max_line_length - len(prefix))
            for i, wrap in enumerate(wrapped):
                wrapped_lines.append(f"{prefix}{wrap}" if i == 0 else f"{' ' * len(prefix)}{wrap}" )
        else:
            wrapped_lines.extend(textwrap.wrap(line,width=max_line_length))
    

    wrapped_text = "\n".join(wrapped_lines)

    display(Markdown(wrapped_text))

  # convert to docx
import markdown2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def strip_html_tags(text):
    import re
    clean = re.compile('<.*?>')
    return re.sub(clean,'',text)

def handle_strong_tags(text):
    segments = text.split("<strong>")
    runs = []
    for segment in segments:
        if "</strong>" in segment:
            bold_text, normal_text =segment.split("</strong>", 1)
            runs.append((bold_text,True))
            runs.append((normal_text, False))
        else:
            runs.append((segment,False))
    return runs

def add_markdown_to_docx(md_text, doc):
    html = markdown2.markdown(md_text)

    lines = html.splitlines()

    for line in lines:
        line = line.strip()

        if line.startswith("<h1>"):
            paragraph = doc.add_heading(level=1)
            paragraph.add_run(strip_html_tags(line))
        elif line.startswith("<h2>"):
            paragraph = doc.add_heading(level=2)
            paragraph.add_run(strip_html_tags(line))
        elif line.startswith("<h3>"):
            paragraph = doc.add_heading(level=3)
            paragraph.add_run(strip_html_tags(line))
        elif "<strong>" in line:
            paragraph = doc.add_paragraph()
            for text, is_bold in handle_strong_tags(line):
                run = paragraph.add_run(strip_html_tags(text))
                run.bold=is_bold
        elif line.startswith("<ul>"):
            continue
        elif line.startswith("</ul>"):
            continue
        elif line.startswith("<li>"):
            paragraph = doc.add_paragraph(style='List Bullet')
            for text, is_bold in handle_strong_tags(line):
                run = paragraph.add_run(strip_html_tags(text))
                run.bold = is_bold
        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run(strip_html_tags(line))

def convert_md_to_docx(md_text, output_filename):
    doc = Document()
    add_markdown_to_docx(md_text, doc)
    doc.save(output_filename)
            
import google.generativeai as genai
import os
api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_key)     

class Agent:
    def __init__(self,name,role):
        self.name=name,
        self.role=role,
        self.model = genai.GenerativeModel('gemini-1.5-flash',system_instruction=role)
    
    def generate_response(self,prompt):
        response = self.model.generate_content(prompt)
        return response.text

cro_role = """ As a Customer Relation Officer, your aim is to gather information about the user's intent and experctations for generating a comperehensive report.

Ask necessary questions to understand the user's spesific intetn and spesific experctations from the report considering the below current conversation.

if your are not sure that you have enough information about the intent and the expectations for generating a comprehensive report keep asking one question at a time to dertermine the uesr's intetn and experctations from the report.

if yout think that your have enough information about both the intent and the expertations for the report, stop asking questions and provide a summary by filling th below format:

SUMMARY:
Intent: [Summarized intent]
Experctations: [List of experctations]

It,s crucial that your summary must begint with SUMMARY keyword in capital letter.

Now, review the current conversation below, and decide whether to ask more questions or not.  

-------------------------------------
The current conversation:
-------------------------------------
"""

author_role = """ As a Report Author, your goal is to write a report in German based on a given user intent and expectations and the improve it according to an editor's revision requests.

First you will be given the user inetent and expectations. Prepare the report in German.  

Then, an Editor will check your report and can submit revision requests to you. When you receive revision requests, revise the latest report accordingly.  
"""

editor_role = """As a Report Editor, your goal is to review the given report based on the given user intent and expectations. 
//////////////////
GENERAL RULES:
/////////////////
Each sections should have at least 3 paragraphs. Each paragraph should have at least 10 sentences.

The report must be to the point and must be concise considering the given user intent and expectations.  

There should be no gramtical or spelling errors.

The conclusion of the report should meet the user's intent and expectations.  

Prepare at least 1 revision request for the report. Your request must be clear, easy to follow, and concise. If poosible, give simple instructions to revise the report step by step.

There should be written in German .   


////////////////
ACTING RULES:
///////////////
First decide if the report needs any revisions according to the user intetn and expectations.  

Then, select one of the below actions:

A. If some revisions are requiered, response with the following format:
***Revisions Required:...***

B. Else, if there is no need for revisons, respons with the following format:
***DONE! Report meets the user's expectations***

"""

class CustomerRelationOfficer(Agent):
    def ger_user_intent(self, max_interaction=3):
        msg= """ Hello! I'm your Customer Relation Officer. Let's discuss your report needs. 
        What's the main topic of your report?
        """
        print("hier1")
        conversation=[f"CRO: {msg}"]
        topic = input(msg)
        print(topic)
        conversation.append(f"User: {topic}")
        for i in range(max_interaction):
            print("hier2")

            prompt = f"""
            -------------------------
            The current conversation:
            -------------------------

            {' '.join(conversation)}
            """

            response = self.generate_response(prompt)
            print_formatted_md(f"CRO: {response}")
            conversation.append((f"CRO: {response}"))

            if response.startswith("SUMMARY:"):
                print("hier3")

                return response
            
            if i < max_interaction-1:
                user_input = input("Your response: ")
                conversation.append(f"User: {user_input}")
                print("heir4")

        summary_prompt = f"""We've reached max interactions. Now, summarize the userÄs intent and expectations based on this conversation:
        {' '.join(conversation)}

        Provide a summary in the format:
        SUMMARY:
        Intent: [Summarized intent]
        Expectations: [List of expectations]

        """
        response = self.generate_response(summary_prompt)
        print_formatted_md(response)
        print("hier5")
        return response
    

    def deliver_report(self,report):
        print("\nYour report is ready. Here it is:")
        print_formatted_md(report)

        report_file = "report.docx"
        convert_md_to_docx(report,report_file)
        print(f"Report saved as {report_file}")

class Author(Agent):
    def write_report(self,user_intent):
        prompt = f"Write a report in German based on this user intent and expectations: {user_intent}"
        return self.generate_response(prompt)
    
    def revise_report(self,original_report, editor_feedback, user_inetnt):
        prompt = f""" Revise this report based on the editor's feedback and the original user intent:

        Original report:
        {original_report}

        Editor's feedback:
        {editor_feedback}

        Original user intent and expectations:
        {user_inetnt}

        Please provide the revised report.
        """
        return self.generate_response(prompt)

class Editor(Agent):
    def review_report(self, report, user_intent):
        prompt = f""" Review this report in German: {report} based
        on this user intent and expectations: {user_intent}
        """
        return self.generate_response(prompt)

def generate_report():
    cro=CustomerRelationOfficer("CRO",cro_role)
    author=Author("Author",author_role)
    editor = Editor("Editor",editor_role)

    user_intent= cro.ger_user_intent()
    print("\n------------CRO-----------------")
    print("\nUser intent and expectations gathered")
    # print(f"\n-----------AUTHOR-----------------")
    # report = author.write_report(user_intent)
    # max_rounds = 3
    # for i in range(max_rounds):

    #     print_formatted_md(f"Draft Report (Round {i+1}): \n{report}\n")
    #     print(f"\n------------EDITOR--------------")
    #     review = editor.review_report(report, user_intent)
    #     print_formatted_md(f"Editor's Review (Round {i+1}):\n{review}\n")

    #     if "meets the user's expectations" in review.lower():
    #         break
    #     elif i<max_rounds:
    #         print(f"\n-----------AUTHOR-----------------")
    #         print("Revising report based on editor's feedback...\n")
    #         report = author.revise_report(report, review, user_intent)
    #     else:
    #         print(f"\n-----------SYSTEM-----------------")
    #         print(f"""Maximum revision rounds ({max_rounds}) reached. 
    #               Proceeding with the current version.""")
    
    # print("\n----------------------CRO------------------")
    # cro.deliver_report(report)

if __name__ == "__main__":
    generate_report()



            